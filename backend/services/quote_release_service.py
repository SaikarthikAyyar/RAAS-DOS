# ====================================
# IMPORTS
# ====================================

import io
import os

from datetime import date
from types import SimpleNamespace

from docx import Document

from backend.models.customer_requests import CustomerRequest
from backend.models.enquiry import Enquiry
from backend.models.machines_pumps import Machine

from backend.repositories.techno_commercial_quote_repository import get_ops_selection, get_quote
from backend.repositories.quote_template_repository import get_active_template
from backend.repositories.quote_release_document_repository import create_quote_release_document

from backend.services.sales_survey_service import get_sales_survey_request

from backend.utils.template_rendering import substitute_tokens

from backend.reporting.quote_release_docx import (
    build_tank_machine_table,
    build_techno_commercial_summary_table,
    add_cover_logo,
    add_header_logo,
    set_document_base_font,
    add_body_paragraph
)


# ====================================
# FALLBACK TEMPLATE BODY
# Never blocks a release - same posture as every other master-lookup
# fallback already established in this codebase (e.g. quote_engine.py's
# accessory-name-mismatch fallback, resolve_machine_rate's service-
# configuration fallback).
# ====================================

FALLBACK_BODY = (
    "Proposal to\nM/s. {customer}\n{site}\n\n"
    "Proposal No: {proposal_no}\nEnquiry reference no: {enquiry_ref}\n"
    "Proposal Submission Date: {proposal_date}\n\n"
    "TANK / MACHINE DETAILS\n{tank_machine_table}\n\n"
    "TECHNO-COMMERCIAL QUOTE SUMMARY\n{commercial_table}\n\n"
    "Valid till: {valid_till}\n\n"
    "Thanks & regards,\nJanyu Technologies Pvt Ltd"
)


# ====================================
# RESOLVE FINAL MACHINE
# Same dual code-or-name tolerance quote_engine.py::resolve_machine_rate
# already relies on - override_machine/recommended_machine is
# inconsistently a code vs. a name today (pre-existing quirk).
# ====================================

def _resolve_final_machine(db, ops):

    if ops is None:
        return None

    identifier = ops.override_machine or ops.recommended_machine

    if not identifier:
        return None

    return (
        db.query(Machine)
        .filter((Machine.code == identifier) | (Machine.name == identifier))
        .first()
    )


# ====================================
# GENERATE
# ====================================

def generate_quote_release_docx(db, quote_id, enquiry_id, generated_by):

    quote = get_quote(db, quote_id)

    if quote is None:
        raise ValueError("Quote not found.")

    ops = get_ops_selection(db, quote.ops_selection_id) if quote.ops_selection_id else None

    sales_survey = (
        get_sales_survey_request(db, ops.sales_survey_id)
        if ops and ops.sales_survey_id
        else None
    )

    customer_request = (
        db.query(CustomerRequest)
        .filter(CustomerRequest.id == quote.customer_request_id)
        .first()
        if quote.customer_request_id
        else None
    )

    enquiry = db.query(Enquiry).filter(Enquiry.id == enquiry_id).first() if enquiry_id else None

    machine = _resolve_final_machine(db, ops)

    template = get_active_template(db)
    body = template.body if template is not None else FALLBACK_BODY

    tokens = {
        "customer": (enquiry.customer_name if enquiry else None)
            or (customer_request.company_name if customer_request else None)
            or "-",
        "site": (customer_request.plant_site_location if customer_request else None) or "-",
        "proposal_no": f"QT-{enquiry_id}-v{quote.revision_number or 1}",
        "enquiry_ref": f"ENQ-{enquiry_id}",
        "proposal_date": date.today().isoformat(),
        "contact_name": (customer_request.contact_person if customer_request else None) or "-",
        "contact_designation": "-",
        "contact_phone": (customer_request.contact_number if customer_request else None) or "-",
        "contact_email": (customer_request.client_contact_email if customer_request else None) or "-",
        "valid_till": quote.valid_till or "-",
        "value": (
            f"{quote.final_approved_value:,.0f}"
            if quote.final_approved_value is not None
            else "-"
        )
    }

    doc = Document()

    set_document_base_font(doc)
    add_header_logo(doc)
    add_cover_logo(doc)

    format_state = {"active": True, "section_number": 0}

    for line in body.split("\n"):

        stripped = line.strip()

        if stripped == "{tank_machine_table}":

            build_tank_machine_table(doc, sales_survey, machine)

        elif stripped == "{commercial_table}":

            build_techno_commercial_summary_table(doc, quote)

        else:

            add_body_paragraph(doc, substitute_tokens(line, tokens), format_state)

    folder = f"backend/uploads/quote_releases/{enquiry_id}"
    os.makedirs(folder, exist_ok=True)

    file_name = f"Quote_ENQ{enquiry_id}_v{quote.revision_number or 1}.docx"
    file_path = f"{folder}/{file_name}"

    doc.save(file_path)

    return create_quote_release_document(
        db,
        quote_id=quote.id,
        enquiry_id=enquiry_id,
        file_name=file_name,
        file_path=file_path,
        generated_by=generated_by
    )


# ====================================
# TEMPLATE PREVIEW
# Lets an admin generate and download a real .docx straight from the
# Quote Templates tab, with no enquiry/quote context required - so
# {tank_machine_table} / {commercial_table} are built from clearly-
# labelled sample data instead of a real Sales Survey/Machine/Quote.
# Not persisted as a QuoteReleaseDocument (it isn't a real release) -
# built and streamed straight from memory.
# ====================================

SAMPLE_TOKENS = {
    "customer": "Sample Customer Pvt Ltd",
    "site": "Sample Plant, Sample City",
    "proposal_no": "QT-SAMPLE-v1",
    "enquiry_ref": "ENQ-SAMPLE",
    "proposal_date": date.today().isoformat(),
    "contact_name": "Sample Contact",
    "contact_designation": "Plant Manager",
    "contact_phone": "+91 98765 43210",
    "contact_email": "sample.contact@example.com",
    "valid_till": "30 days from proposal date",
    "value": "25,00,000"
}

SAMPLE_SALES_SURVEY = SimpleNamespace(
    tank_type="Cuboidal",
    tank_length=15,
    tank_width=10,
    tank_depth=6,
    estimated_volume=900,
    material_category="Heavy Sludge",
    job_type="Tank Cleaning"
)

SAMPLE_MACHINE = SimpleNamespace(
    name="Sample Machine XYZ",
    power_type="Diesel",
    base_output_per_day=40,
    base_output_basis="m3 sludge per shift",
    max_vertical_lift=6,
    hazard_rating="Standard"
)

SAMPLE_QUOTE = SimpleNamespace(
    mobilisation_cost_min=175000, mobilisation_cost_max=175000,
    setup_cost_min=125000, setup_cost_max=125000,
    execution_cost_min=375000, execution_cost_max=450000,
    pump_addon_cost_min=10000, pump_addon_cost_max=15000,
    documentation_buffer=15000,
    access_support_buffer=25000,
    overhead_cost_min=108750, overhead_cost_max=127500,
    contingency_cost_min=72500, contingency_cost_max=85000,
    margin_value_min=226563, margin_value_max=265625,
    dewatering_addon_min=0, dewatering_addon_max=50000,
    combined_budgetary_value_min=1132813, combined_budgetary_value_max=1333125,
    final_approved_value=1200000
)


def generate_quote_template_preview_docx(template):

    doc = Document()

    set_document_base_font(doc)
    add_header_logo(doc)
    add_cover_logo(doc)

    body = template.body or ""

    if not body.strip():

        doc.add_paragraph("This template has no body yet - add content via Edit first.")

    format_state = {"active": True, "section_number": 0}

    for line in body.split("\n"):

        stripped = line.strip()

        if stripped == "{tank_machine_table}":

            build_tank_machine_table(doc, SAMPLE_SALES_SURVEY, SAMPLE_MACHINE)

        elif stripped == "{commercial_table}":

            build_techno_commercial_summary_table(doc, SAMPLE_QUOTE)

        else:

            add_body_paragraph(doc, substitute_tokens(line, SAMPLE_TOKENS), format_state)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
