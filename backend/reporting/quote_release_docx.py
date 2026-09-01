# ====================================
# QUOTE RELEASE DOCX - FORMATTING + DYNAMIC TABLES
#
# The two segments a Quote Template can't reasonably hand-type per
# quote, because they're relational (tank/machine specs) or derived
# (the commercial rate) rather than boilerplate. Everything else in
# the document is plain admin-edited text, substituted via the same
# {token} mechanism Email Templates already uses - but that plain
# text still needs real document formatting (fonts, headings, spacing)
# to read as a professional proposal rather than a raw text dump.
# This file's classifier is what does that, styled after a real
# reference proposal (unpacked and inspected directly: Times New
# Roman throughout, centered ALL-CAPS section headings, bold inline
# labels before a colon, justified body text, a centered cover block).
# ====================================

import os

from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ====================================
# LOGO
# Two places, matching the real reference document's own actual dual
# treatment (confirmed by unpacking its raw XML): a large logo on the
# cover page itself (an inline image at the very top of the document
# body, before any text), plus a smaller logo running in the top-right
# of the page header from page 2 onward - never in the header on the
# cover page itself, since Word's own reference document keeps its
# "first page" header genuinely blank (a separate header part from its
# "default" header) and relies on the body-level cover image instead.
# python-docx exposes that same first-page/default split as
# Section.different_first_page_header_footer + Section.first_page_header.
# ====================================

LOGO_PATH = os.path.join("frontend", "src", "assets", "JanyutechLogo.jpg")


def add_cover_logo(doc):

    if not os.path.exists(LOGO_PATH):
        return

    paragraph = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    run.add_picture(LOGO_PATH, width=Inches(4.0))


def add_header_logo(doc):

    if not os.path.exists(LOGO_PATH):
        return

    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False

    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = paragraph.add_run()
    run.add_picture(LOGO_PATH, width=Inches(1.6))


# ====================================
# DOCUMENT-WIDE BASE FONT
# Times New Roman throughout, matching the reference proposal - set
# on the Normal style so every paragraph inherits it unless overridden.
# ====================================

def set_document_base_font(doc):

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)


# ====================================
# BODY PARAGRAPH CLASSIFIER
# The template body is plain admin-edited text (one line = one
# paragraph) - this is what turns that flat text into something that
# reads as a real proposal: ALL-CAPS lines become centered underlined
# section headings (auto-numbered "1.", "2.", ...), short colon-
# terminated lines become bold left sub-headings, an inline
# "Label: value" prefix gets bolded in place, lines starting with
# "- " become real Word bullet-list items (the marker stripped, not
# left as literal text) - matching the reference document's own real
# dash-bulleted lists (sector list, JanyuTech/Client responsibility
# lists), and everything before the first section heading (the cover
# block - "Proposal to" / customer / site / proposal metadata /
# letterhead) is centered and stays on its own page, matching the
# reference document's own cover-page treatment.
#
# Every heading starts a fresh page by default - direct instruction:
# Introduction gets a full page to itself, the machine section and
# Commercial Proposal must each start a new page. The one deliberate
# exception is "same page as the section before it" (Tank / Machine
# Details staying with Project Overview), opted into per-heading by
# prefixing the template line with "^" (stripped before rendering,
# never shown) - explicit per-section control instead of a single
# blanket rule, since different sections genuinely need different
# behavior here.
# ====================================

def _is_all_caps_heading(text):

    if not text or len(text) > 60:
        return False

    if not any(c.isalpha() for c in text):
        return False

    if any(c.isdigit() for c in text):
        return False

    return text.isupper()


def _is_short_subheading(text):

    return bool(text) and len(text) <= 50 and text.endswith(":")


def _is_bullet_line(text):

    return text.startswith("- ") and len(text) > 2


# ====================================
# SECTION HEADING
# Shared by add_body_paragraph's own ALL-CAPS detection AND
# build_machine_section below, so a heading whose TEXT is dynamic
# (the machine's real name, not a fixed label) still renders with
# exactly the same numbering/styling/page-break convention as every
# other section heading in the document - one real mechanism, not two
# diverging ones.
# ====================================

def add_section_heading(doc, format_state, heading_text, force_page_break=True):

    format_state["active"] = False
    format_state["section_number"] += 1

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if force_page_break:
        paragraph.paragraph_format.page_break_before = True

    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)

    run = paragraph.add_run(f"{format_state['section_number']}. {heading_text}")
    run.bold = True
    run.underline = True
    run.font.size = Pt(13)

    return paragraph


def add_body_paragraph(doc, raw_line, format_state):

    stripped = raw_line.strip()

    if not stripped:
        doc.add_paragraph("")
        return

    if stripped.startswith("^") and _is_all_caps_heading(stripped[1:].strip()):

        add_section_heading(doc, format_state, stripped[1:].strip(), force_page_break=False)
        return

    if _is_all_caps_heading(stripped):

        add_section_heading(doc, format_state, stripped, force_page_break=True)
        return

    if _is_bullet_line(stripped):

        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(stripped[2:])

        return

    if not format_state["active"] and _is_short_subheading(stripped):

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(2)

        run = paragraph.add_run(stripped)
        run.bold = True

        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if format_state["active"] else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(4)

    colon_index = stripped.find(": ")

    if stripped.endswith(":") and colon_index == -1:

        run = paragraph.add_run(stripped)
        run.bold = True

    elif 0 < colon_index <= 45:

        label_run = paragraph.add_run(stripped[:colon_index + 1])
        label_run.bold = True

        paragraph.add_run(stripped[colon_index + 1:])

    else:

        paragraph.add_run(stripped)


# ====================================
# TANK / MACHINE DETAILS TABLE
# ====================================

def build_tank_machine_table(doc, sales_survey, machine):

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"

    header = table.rows[0].cells
    header[0].text = "Parameter"
    header[1].text = "Value"

    for cell in header:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    rows = []

    if sales_survey is not None:

        rows.append(("Tank Type", sales_survey.tank_type or "-"))
        rows.append(("Tank Length / Diameter (m)", sales_survey.tank_length or "-"))
        rows.append(("Tank Width (m)", sales_survey.tank_width or "-"))
        rows.append(("Tank Depth (m)", sales_survey.tank_depth or "-"))
        rows.append(("Estimated Volume (m3)", sales_survey.estimated_volume or "-"))

    if machine is not None:

        rows.append(("Recommended Machine", machine.name or "-"))
        rows.append(("Power Type", machine.power_type or "-"))

        base_output = (
            f"{machine.base_output_per_day} {machine.base_output_basis}"
            if machine.base_output_per_day
            else "-"
        )
        rows.append(("Base Output Capacity", base_output))
        rows.append(("Max Vertical Lift (m)", machine.max_vertical_lift or "-"))
        rows.append(("Hazard Rating", machine.hazard_rating or "-"))

    if not rows:
        rows.append(("-", "No tank or machine data available for this enquiry yet."))

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    return table


# ====================================
# COMMERCIAL RATE TABLE
# Full techno-commercial rate breakdown - every line item that sums
# to the combined budgetary value (mobilisation, setup, execution,
# pump addon, documentation/access-support buffers, overhead,
# contingency, margin, dewatering add-on), plus the range and the
# final approved value - matching exactly what's already shown on the
# Commercial Approval / Quote & Commercial tabs' "Finalized quote
# lines" table, per direct instruction.
# ====================================

# ====================================
# MACHINE SECTION (fully dynamic - heading AND content)
# The reference proposal's own "Section 4" is a real, standalone
# numbered section headed by the machine's own name (e.g.
# "4.VARAHA - SCH Sump Cleaning Robot") - not a sub-heading nested
# inside the tank/machine details section. This reproduces that
# exact structure: a real section heading (participating in the same
# auto-numbering as every other section, via add_section_heading) whose
# TEXT is the resolved machine's real name, starting a fresh page, with
# everything below it built the same way - every real fact mapped onto
# a real, structured Machine Spec field, not hardcoded to any one
# machine's copy:
#
#   reference phrase                          -> real field
#   "4.VARAHA - SCH Sump Cleaning Robot"      -> the section heading itself = machine.name
#   "100% Hydraulically operated..."          -> machine.power_type
#   "Suitable for ATEX Zone 0 area"           -> machine.hazard_rating
#   "Fits through a man entry hole, Ø 600mm"  -> machine.minimum_width / minimum_height
#     (the machine's own real access-opening requirement)
#   "Robotic arm equipped with camera and
#    mount for Water jet nozzle"              -> this job's ops_selection.accessories_plan
#     (the Deployment Plan's real Yes/No checklist - what's actually
#      equipped for THIS job, not the machine's generic default list)
#   "Controlling from a safe distance..."     -> machine.vehicle / vehicle_payload
#     (how it's actually deployed to site)
#   opening descriptive paragraph             -> machine.description
#     (a real, admin-editable free-text field on the Machine Spec,
#      already populated for real machines, e.g. "Diesel operated
#      floating sludge removal platform." for Matsya D)
#
# Every bullet is conditional on the underlying field actually being
# set - a machine with no crane requirement, no rated vertical lift,
# etc. simply omits that line rather than showing a fabricated value.
# The reference document's own actual product photo is deliberately
# NOT reproduced here (confirmed with the user: images are added
# manually afterward, not generated) - this only produces the real,
# structured facts a photo can't carry.
# ====================================

def build_machine_section(doc, format_state, machine, accessories_plan):

    if machine is None:

        add_section_heading(doc, format_state, "MACHINE DETAILS", force_page_break=True)

        paragraph = doc.add_paragraph()
        paragraph.add_run(
            "No machine has been finalized for this enquiry yet."
        )
        return

    add_section_heading(doc, format_state, machine.name, force_page_break=True)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if machine.description:
        intro.add_run(machine.description)
    else:
        intro.add_run(
            f"{machine.name} is the machine finalized for this job, "
            f"matched against the survey's real requirements via Ops Selection."
        )

    features_label = doc.add_paragraph()
    features_label.paragraph_format.space_before = Pt(6)
    features_label.paragraph_format.space_after = Pt(2)
    features_label_run = features_label.add_run("Key Features:")
    features_label_run.bold = True

    bullets = []

    if machine.power_type:
        bullets.append(f"{machine.power_type} operated system.")

    if machine.hazard_rating and machine.hazard_rating != "Standard":
        bullets.append(f"Suitable for {machine.hazard_rating} hazardous area rating.")

    if machine.minimum_width and machine.minimum_height:
        bullets.append(
            f"Fits through an access opening of {machine.minimum_width}mm (W) "
            f"x {machine.minimum_height}mm (L) or larger."
        )
    elif machine.minimum_width:
        bullets.append(f"Requires a minimum access opening width of {machine.minimum_width}mm.")
    elif machine.minimum_height:
        bullets.append(f"Requires a minimum access opening length of {machine.minimum_height}mm.")

    if machine.max_vertical_lift:
        bullets.append(f"Vertical lift/reach capability of up to {machine.max_vertical_lift}m.")

    if machine.crane_required == "Yes":
        bullets.append("Crane support required for deployment/retrieval.")

    if machine.base_output_per_day:
        basis = f" {machine.base_output_basis}" if machine.base_output_basis else ""
        bullets.append(f"Rated output of {machine.base_output_per_day}{basis}.")

    equipped = [
        item.get("name")
        for item in (accessories_plan or [])
        if isinstance(item, dict) and item.get("needed") == "Yes" and item.get("name")
    ]

    if equipped:
        bullets.append(f"Equipped for this job with: {', '.join(equipped)}.")

    if machine.vehicle:
        payload = f" ({machine.vehicle_payload} payload)" if machine.vehicle_payload else ""
        bullets.append(f"Deployed via {machine.vehicle}{payload}.")

    if not bullets:
        bullets.append(
            "Detailed technical specifications for this machine are on file with Janyu Technologies."
        )

    for bullet_text in bullets:

        bullet_paragraph = doc.add_paragraph(style="List Bullet")
        bullet_paragraph.add_run(bullet_text)

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disclaimer.paragraph_format.space_before = Pt(6)

    disclaimer_run = disclaimer.add_run(
        "The items and configuration described above are indicative and may be "
        "adjusted as per final technical assessment at site."
    )
    disclaimer_run.italic = True


def build_techno_commercial_summary_table(doc, quote):

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"

    header = table.rows[0].cells
    header[0].text = "Line"
    header[1].text = "Min (INR)"
    header[2].text = "Max (INR)"

    for cell in header:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    lines = [
        ("Mobilisation", quote.mobilisation_cost_min, quote.mobilisation_cost_max),
        ("Setup / Access", quote.setup_cost_min, quote.setup_cost_max),
        ("Execution (Machine)", quote.execution_cost_min, quote.execution_cost_max),
        ("Pump Addon", quote.pump_addon_cost_min, quote.pump_addon_cost_max),
        ("Documentation Buffer", quote.documentation_buffer, quote.documentation_buffer),
        ("Access Support Buffer", quote.access_support_buffer, quote.access_support_buffer),
        ("Overhead", quote.overhead_cost_min, quote.overhead_cost_max),
        ("Contingency", quote.contingency_cost_min, quote.contingency_cost_max),
        ("Margin", quote.margin_value_min, quote.margin_value_max),
        ("Dewatering Add-on", quote.dewatering_addon_min, quote.dewatering_addon_max)
    ]

    for label, min_value, max_value in lines:

        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = f"{min_value:,.0f}" if min_value is not None else "-"
        cells[2].text = f"{max_value:,.0f}" if max_value is not None else "-"

    range_row = table.add_row().cells
    range_row[0].text = "Combined Budgetary Value (Range)"
    range_row[1].text = f"{quote.combined_budgetary_value_min:,.0f}" if quote.combined_budgetary_value_min is not None else "-"
    range_row[2].text = f"{quote.combined_budgetary_value_max:,.0f}" if quote.combined_budgetary_value_max is not None else "-"

    for cell in range_row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    if quote.final_approved_value is not None:

        final_row = table.add_row().cells
        final_row[0].text = "Final Approved Value"

        merged_value_cell = final_row[1].merge(final_row[2])
        merged_value_cell.text = f"{quote.final_approved_value:,.0f}"

        for paragraph in final_row[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

        for paragraph in merged_value_cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    return table
